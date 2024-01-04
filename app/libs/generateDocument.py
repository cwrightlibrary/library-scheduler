from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from os.path import dirname, join, realpath

# Guide for how the date, leave, and programsMeetings should be formatted:
# date = ["Thursday, January 4, 2024", "thursday_01_04_2024"]
# leave = [[["yami"], "12:00", "2:00"], [["chris"], 0, 0]]
# programsMeetings = [[["lea"], "9:00", "11:00", "Storytime"], [["anthony"], "6:00", "8:00", "Astronomy"]]

class Schedule():
	def __init__(self, date, leave, programsMeetings):
		# Initialize time slots and create class variables for leave/programs/meetings
		self.date = date
		self.leave = leave
		self.programsMeetings = programsMeetings
		self.assignTimeSlots()

		# Initialize the document
		self.document = Document()
		sections = self.document.sections
		for section in sections:
			section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
			newWidth, newHeight = section.page_height, section.page_width
			section.page_width, section.page_height = newWidth, newHeight
			section.orientation = WD_ORIENT.LANDSCAPE
		
		# Adjust Word headings/paragraph theming
		dateStyle = self.document.styles["Heading 1"]
		dateStyle.paragraph_format.space_before = dateStyle.paragraph_format.space_after = 1
		dateStyle.font.name, dateStyle.font.size = "Aptos Display", Pt(16)

		textStyle = self.document.styles["Normal"]
		textStyle.paragraph_format.space_before = textStyle.paragraph_format.space_after = 1
		textStyle.font.name, textStyle.font.size = "Aptos", Pt(11)

		# Adding HEADER INFORMATION to the document
		# Add the date to the document
		printDate = self.document.add_paragraph()
		printDate.style = self.document.styles["Heading 1"]
		printDate.alignment = 1
		self.addDate = printDate.add_run(date[0])
		self.addDate.bold = True
		self.addDate.font.color.rgb = RGBColor(0, 0, 0)
		self.addDate.font.name = "Aptos Display"



	def assignTimeSlots(self):
		if "sunday" in self.date[0].lower():
			self.timeSlots = ["two_three", "three_four", "four_five", "five_six"]
			self.timeCompare = [[14, 15], [15, 16], [16, 17], [17, 18]]
		elif "friday" in self.date[0].lower() or "saturday" in self.date[0].lower():
			self.timeSlots = ["nine_eleven", "eleven_twelve", "twelve_one", "one_two", "two_four", "four_six"]
			self.timeCompare = [[9, 11], [11, 12], [12, 13], [13, 14], [14, 16], [16, 18]]
		else:
			self.timeSlots = ["nineEleven", "elevenOne", "oneTwo", "twoFour", "fourSix", "sixEight"]
			self.timeCompare = [[9, 11], [11, 13], [13, 14], [14, 16], [16, 18], [18, 20]]
	
	def setCellBackground(self, cell, fill):
		from docx.oxml.shared import qn
		from docx.oxml.xmlchemy import OxmlElement
		cellProperties = cell._element.tcPr
		try:
			cellShading = cellProperties.xpath("w:shd")[0]
		except IndexError:
			cellShading = OxmlElement("w:shd")
		if fill:
			cellShading.set(qn("w:fill"), fill)
		cellProperties.append(cellShading)
	
	def exportDocument(self, openFile):
		filename = join(dirname(realpath(__file__)), self.date[1] + ".docx")
		self.document.save(join(filename))

		if openFile:
			from os import name
			if name == "posix":
				from subprocess import call
				call(("open", join(filename)))
			elif name == "nt":
				from os import startfile
				startfile(join(filename))


date = ["Thursday, January 4, 2024", "thursday_01_04_2024"]
leave = [[["yami"], "12:00", "2:00"], [["chris"], 0, 0]]
programsMeetings = [[["lea"], "9:00", "11:00", "Storytime"], [["anthony"], "6:00", "8:00", "Astronomy"]]
schedule = Schedule(date, leave, programsMeetings)
schedule.exportDocument(openFile=False)