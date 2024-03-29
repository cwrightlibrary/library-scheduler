from csv_generators import staff, template, location, convert_time, TEMPLATE
from datetime import *
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from os.path import dirname, join, realpath
from prettytable import PrettyTable


class Schedule:
    def __init__(self, date, adjustments):
        # Get the library branch's name
        self.branch_name = location.location_info["location-name"]
        # Initialize a dictionary for the branch's hours
        self.branch_hours = {}
        # Initialize a list for the branch's floor locations
        self.floor_locations = []
        self.floor_categories = [
            "workroom",
            "front desk",
            "computer desk",
            "programs & projects",
        ]
        # Set the date to the argument
        self.date = date
        # Get the weekday
        self.weekday = datetime.strptime(self.date, "%B %d, %Y").strftime("%A")
        # Set the schedule adjustments to the argument
        self.leave = adjustments[0]
        self.programs = adjustments[1]
        # Create a table for the schedule information
        self.info_header = ["who works today", "lunch breaks", "schedule changes"]

        self.get_hours_locations()
        self.create_header_compare()

        self.get_weekday()
        self.get_daily_hours()
        
        self.get_leave_and_programs()

        self.add_to_schedule()
        self.pretty_print()

    def get_hours_locations(self):
        # Loop through the imported location csv's info dictionary
        for k, v in location.location_info.items():
            # Ignore the branch name and floor locations
            if "hours" in k and v.lower() != "closed":
                day = k.capitalize()
                day = day.replace("-hours", "")
                # Split the formatted "9:00am-8:00pm" string into ["9:00am", "8:00pm"]
                hours = v.split("-")
                # Convert the time to 24 hour integers
                hours[0] = int(convert_time(hours[0], to_24=True))
                hours[1] = int(convert_time(hours[1], to_24=True))
                # Assign dictionary values for each day
                self.branch_hours[day] = hours
            if k == "floor-locations":
                floor_locations = v.split("/")
                for loc in floor_locations:
                    self.floor_locations.append(loc.replace("-", " "))

        # Remove the "workroom" floor category if there's no pickup window
        if not "pickup window" in self.floor_locations:
            self.floor_categories.pop(0)

    def create_header_compare(self):
        for k, v in self.branch_hours.items():
            if self.weekday == k:
                self.today_hours = v
                # If the hours are 2-6, set the appropriate header/compare_time
                if v == [1400, 1800]:
                    self.word_doc = join(
                        dirname(realpath(__file__)),
                        "templates",
                        "two_six_template.docx",
                    )
                    self.template_header = ["", "2 - 3", "3 - 4", "4 - 5", "5 - 6"]
                    self.compare_time = [
                        [1400, 1500],
                        [1500, 1600],
                        [1600, 1700],
                        [1700, 1800],
                    ]
                # If the hours are 9-6, set the appropriate header/compare_time
                elif v == [900, 1800]:
                    self.word_doc = join(
                        dirname(realpath(__file__)),
                        "templates",
                        "nine_six_template.docx",
                    )
                    self.template_header = [
                        "",
                        "9 - 11",
                        "11 - 12",
                        "12 - 1",
                        "1 - 2",
                        "2 - 4",
                        "4 - 6",
                    ]
                    self.compare_time = [
                        [900, 1100],
                        [1100, 1200],
                        [1200, 1300],
                        [1300, 1400],
                        [1400, 1600],
                        [1600, 1800],
                    ]
                # If the hours are 9-8, set the appropriate header/compare_time
                elif v == [900, 2000]:
                    self.word_doc = join(
                        dirname(realpath(__file__)),
                        "templates",
                        "nine_eight_template.docx",
                    )
                    self.template_header = [
                        "",
                        "9 - 11",
                        "11 - 1",
                        "1 - 2",
                        "2 - 4",
                        "4 - 6",
                        "6 - 8",
                    ]
                    self.compare_time = [
                        [900, 1100],
                        [1100, 1300],
                        [1300, 1400],
                        [1400, 1600],
                        [1600, 1800],
                        [1800, 2000],
                    ]

        # Create PrettyTables for the template and information
        self.print_template = PrettyTable(self.template_header)
        self.print_info = PrettyTable(self.info_header)

    def get_leave_and_programs(self):
        for l in self.leave:
            for name in l[0]:
                for employee in staff:
                    if employee["name-shorthand"] == name:
                        if 0 in l[1]:
                            hour1, hour2 = "all day", "all day"
                        else:
                            hour1, hour2 = convert_time(l[1][0]), convert_time(l[1][1])
                        employee["leave"] = [hour1, hour2]
        for p in self.programs:
            for name in p[0]:
                for employee in staff:
                    if employee["name-shorthand"] == name:
                        hour1, hour2 = convert_time(p[1][0]), convert_time(p[1][1])
                        employee["program"] = [hour1, hour2]
        self.leave_print = self.leave
        self.programs_print = self.programs
        for l in range(len(self.leave_print)):
            for name in range(len(self.leave_print[l][0])):
                for employee in staff:
                    if employee["name-shorthand"] == self.leave_print[l][0][name]:
                        self.leave_print[l][0][name] = employee["name"].split()[0]
        for p in range(len(self.programs_print)):
            for name in range(len(self.programs_print[p][0])):
                for employee in staff:
                    if employee["name-shorthand"] == self.programs_print[p][0][name]:
                        self.programs_print[p][0][name] = employee["initials"]
            self.programs_print[p][2] = self.programs_print[p][2].title()

        leave_string = "leave\n"
        programs_string = "programs\n"
        for info in self.leave_print:
            if 0 in info[1]:
                leave_string += ", ".join(info[0]) + "\n"
            else:
                leave_string += (
                    "-".join(info[1])
                    .replace("am", "")
                    .replace("pm", "")
                    .replace(":00", "")
                    + ": "
                )
                leave_string += ", ".join(info[0]) + "\n"
        for info in self.programs_print:
            programs_string += (
                "-".join(info[1]).replace("am", "").replace("pm", "").replace(":00", "")
                + ": "
                + info[2]
                + " "
            )
            programs_string += "(" + ", ".join(info[0]) + ")\n"
        self.print_info.add_row([self.hours_string, self.break_string, leave_string + "\n" + programs_string])

    def get_weekday(self):
        self.current_date = []

        def add_to_day(day, amount):
            return datetime.strftime(
                datetime.strptime(day, "%B %d, %Y") + timedelta(days=amount),
                "%B %d, %Y",
            )

        (
            friday1a,
            friday1b,
            friday2a,
            friday2b,
            friday3a,
            friday3b,
            saturday1a,
            saturday1b,
            saturday2a,
            saturday2b,
            saturday3a,
            saturday3b,
            sunday1,
            sunday2,
            sunday3,
        ) = (
            [TEMPLATE.a_weekend_1],
            [TEMPLATE.a_weekend_2],
            [TEMPLATE.b_weekend_1],
            [TEMPLATE.b_weekend_2],
            [TEMPLATE.c_weekend_1],
            [TEMPLATE.c_weekend_2],
            [add_to_day(TEMPLATE.a_weekend_1, 1)],
            [add_to_day(TEMPLATE.a_weekend_2, 1)],
            [add_to_day(TEMPLATE.b_weekend_1, 1)],
            [add_to_day(TEMPLATE.b_weekend_2, 1)],
            [add_to_day(TEMPLATE.c_weekend_1, 1)],
            [add_to_day(TEMPLATE.c_weekend_2, 1)],
            [add_to_day(TEMPLATE.a_weekend_1, 2)],
            [add_to_day(TEMPLATE.b_weekend_1, 2)],
            [add_to_day(TEMPLATE.c_weekend_1, 2)],
        )
        weekends = [
            friday1a,
            friday1b,
            friday2a,
            friday2b,
            friday3a,
            friday3b,
            saturday1a,
            saturday1b,
            saturday2a,
            saturday2b,
            saturday3a,
            saturday3b,
            sunday1,
            sunday2,
            sunday3,
        ]
        for day in weekends:
            for amt in range(0, 500):
                three_weeks = amt * 21
                day.append(add_to_day(day[0], three_weeks))
        for day in range(len(weekends)):
            if self.date in weekends[day]:
                if day == 0:
                    self.current_date = ["friday-hours", 0, 0]
                if day == 1:
                    self.current_date = ["friday-hours", 0, 1]
                if day == 2:
                    self.current_date = ["friday-hours", 1, 0]
                if day == 3:
                    self.current_date = ["friday-hours", 1, 1]
                if day == 4:
                    self.current_date = ["friday-hours", 2, 0]
                if day == 5:
                    self.current_date = ["friday-hours", 2, 1]
                if day == 6:
                    self.current_date = ["saturday-hours", 0, 0]
                if day == 7:
                    self.current_date = ["saturday-hours", 0, 1]
                if day == 8:
                    self.current_date = ["saturday-hours", 1, 0]
                if day == 9:
                    self.current_date = ["saturday-hours", 1, 1]
                if day == 10:
                    self.current_date = ["saturday-hours", 2, 0]
                if day == 11:
                    self.current_date = ["saturday-hours", 2, 1]
                if day == 12:
                    self.current_date = ["sunday-hours", 0]
                if day == 13:
                    self.current_date = ["sunday-hours", 1]
                if day == 14:
                    self.current_date = ["sunday-hours", 2]
        if self.weekday.lower() not in ["friday", "saturday", "sunday"]:
            self.current_date = self.weekday.lower() + "-hours"
            self.current_date_string = self.weekday.lower()
        if isinstance(self.current_date, list):
            if len(self.current_date) == 2:
                if self.current_date[1] == 0: m1 = "1"
                if self.current_date[1] == 1: m1 = "2"
                if self.current_date[1] == 2: m1 = "3"
                self.current_date_string = self.current_date[0].replace("-hours", "") + m1
            else:
                if self.current_date[1] == 0: m1 = "1"
                if self.current_date[1] == 1: m1 = "2"
                if self.current_date[1] == 2: m1 = "3"
                if self.current_date[2] == 0: m2 = "a"
                if self.current_date[2] == 1: m2 = "b"
                self.current_date_string = self.current_date[0].replace("-hours", "") + m1 + m2

    def get_daily_hours(self):
        hours = []
        combined_hours = []
        self.hours_string = ""
        self.break_string = ""
        for employee in staff:
            if self.weekday.lower() in ["friday", "saturday"]:
                employee_hours = employee[self.current_date[0]][self.current_date[1]][
                    self.current_date[2]
                ]
            elif self.weekday.lower() == "sunday":
                employee_hours = employee[self.current_date[0]][self.current_date[1]]
            else:
                employee_hours = employee[self.current_date]
            if employee_hours != "Off" and employee_hours not in hours:
                hours.append(employee_hours)
        hours = sorted(hours)
        for hour in hours:
            print_hours = [convert_time(hour[0], to_24=False) + "-" + convert_time(hour[1], to_24=False) + ":\n"]
            names = []
            for employee in staff:
                if self.weekday.lower() in ["friday", "saturday"]:
                    employee_hours = employee[self.current_date[0]][self.current_date[1]][
                        self.current_date[2]
                    ]
                elif self.weekday.lower() == "sunday":
                    employee_hours = employee[self.current_date[0]][self.current_date[1]]
                else:
                    employee_hours = employee[self.current_date]
                if hour[0] == employee_hours[0] and hour[1] == employee_hours[1]:
                    names.append(employee["name"].split()[0])
            print_hours.append(names)
            combined_hours.append(print_hours)
        for hour in combined_hours:
            self.hours_string += hour[0] + ", ".join(hour[1]) + "\n\n"

        hours = []
        combined_hours = []
        for employee in staff:
            if self.weekday.lower() in ["friday", "saturday"]:
                employee_hours = employee[self.current_date[0]][self.current_date[1]][
                    self.current_date[2]
                ]
                employee_break = employee[self.current_date[0].replace("hours", "break")][self.current_date[1]][
                    self.current_date[2]
                ]
            elif self.weekday.lower() == "sunday":
                employee_hours = employee[self.current_date[0]][self.current_date[1]]
                employee_break = employee[self.current_date[0].replace("hours", "break")][self.current_date[1]]
            else:
                employee_hours = employee[self.current_date]
                employee_break = employee[self.current_date.replace("hours", "break")]
            if employee_hours != "Off" and int(employee_hours[1]) - int(employee_hours[0]) >= 600 and employee_break != "None":
                break_hours = [employee_break[0], str(int(employee_break[0]) + 100)]
                if break_hours not in hours:
                    hours.append(break_hours)
        hours = sorted(hours)
        for hour in hours:
            print_hours = [convert_time(hour[0], to_24=False) + "-" + convert_time(hour[1], to_24=False) + ":\n"]
            names = []
            for employee in staff:
                if self.weekday.lower() in ["friday", "saturday"]:
                    employee_hours = employee[self.current_date[0]][self.current_date[1]][
                        self.current_date[2]
                    ]
                    employee_break = employee[self.current_date[0].replace("hours", "break")][self.current_date[1]][
                        self.current_date[2]
                    ]
                elif self.weekday.lower() == "sunday":
                    employee_hours = employee[self.current_date[0]][self.current_date[1]]
                    employee_break = employee[self.current_date[0].replace("hours", "break")][self.current_date[1]]
                else:
                    employee_hours = employee[self.current_date]
                    employee_break = employee[self.current_date.replace("hours", "break")]
                if employee_hours != "Off" and int(employee_hours[1]) - int(employee_hours[0]) >= 600 and employee_break != "None":
                    break_hours = [employee_break[0], str(int(employee_break[0]) + 100)]
                elif employee_hours != "Off" and int(employee_hours[1]) - int(employee_hours[0]) < 600 and employee_break != "None":
                    break_hours = [employee_break[0], str(int(employee_break[0]) + 30)]
                if hour[0] == break_hours[0] and hour[1] == break_hours[1] and hour[0] == employee_break[0]:
                    names.append(employee["name"].split()[0])
            print_hours.append(names)
            combined_hours.append(print_hours)
        for hour in combined_hours:
            self.break_string += hour[0] + ", ".join(hour[1]) + "\n\n"
        
    def add_to_schedule(self):
        # Create list for all available location/hour slots
        self.template = []
        all_locs = [self.current_date_string + "PUW", self.current_date_string + "FL", self.current_date_string + "SP1a", self.current_date_string + "SP1b", self.current_date_string + "SP2a", self.current_date_string + "SP2b"]
        name_locs = ["pickup window", "floor lead", "service point 1", "service point 1", "service point 2", "service point 2"]
        
        for k, v in template.items():
            for loc in range(len(all_locs)):
                if k == all_locs[loc]:
                    temp = []
                    temp.append(name_locs[loc])
                    for name in v:
                        if "/" in name:
                            name = name.split("/")
                            if "*" in name[0]:
                                name_string = ""
                                name[0] = name[0].split("*")
                                # TODO If the name starts with *
                                print(name[0])
                                for employee in staff:
                                    if employee["name-shorthand"] == name[0][0]:
                                        name_string = employee["name"].split()[0] + " 'til " + name[0][1].replace("am", "").replace("pm", "")
                                    if employee["name-shorthand"] == name[1]:
                                        name = name_string + "\n" + employee["name"].split()[0]
                            temp.append(name)
                        else:
                            if name == None or name == "" or name == "none":
                                name = ""
                                temp.append(name)
                            else:
                                for employee in staff:
                                    if employee["name-shorthand"] == name:
                                        if employee["leave"] != []:
                                            idx = v.index(name)
                                            if "all day" in employee["leave"]:
                                                temp.append("")
                                            else:
                                                if not self.compare_hours(int(employee["leave"][0]), int(employee["leave"][1]), int(self.compare_time[idx][0]), int(self.compare_time[idx][1])):
                                                    temp.append("")
                                                else:
                                                    temp.append(employee["name"].split()[0])
                                        else:
                                            temp.append(employee["name"].split()[0])
                    print(temp)
                    self.template.append(temp)
        
        # TEMPORARY FOR FUNCTIONALITY
        # self.template = []
        # for loc in range(len(self.floor_locations)):
        #     temp = []
        #     temp.append(self.floor_locations[loc])
        #     for hour in range(len(self.template_header) - 1):
        #         temp.append("")
        #     self.template.append(temp)

        for row in range(len(self.template)):
            self.print_template.add_row(self.template[row], divider=True)

    def pretty_print(self):
        # Set the padding width for the PrettyTables
        self.print_info.padding_width = 20
        self.print_template.padding_width = 10
        # Get the width of the info table
        table_width = len(self.print_info.get_string().splitlines()[0])
        # Strings to center the weekday/date over the PrettyTable
        weekday_centered = (
            " " * int((table_width / 2) - int(len(self.weekday) / 2)) + self.weekday
        )
        date_centered = (
            " " * int((table_width / 2) - int(len(self.date.replace(",", "")) / 2))
            + self.date
        )

        # Print all
        print(weekday_centered)
        print(date_centered)
        print(self.print_info)
        print(self.print_template)

    def change_doc(self):
        black = RGBColor(20, 20, 20)
        aptos = "Aptos"
        aptos_display = "Aptos Display"

        self.document = Document(self.word_doc)

        self.document.paragraphs[0].text = self.weekday + ", " + self.date
        self.document.paragraphs[0].runs[0].font.name = aptos_display
        self.document.paragraphs[0].runs[0].font.color.rgb = black

        self.document.save(join(dirname(realpath(__file__)), "test.docx"))
    
    def compare_hours(self, hour1, hour2, comp1, comp2):
        if ((hour1 >= comp1 and hour2 <= comp2)
        or (hour1 < comp1 and hour2 >= comp2)
        or (hour1 < comp1 and hour2 >= comp1 and hour2 <= comp2)):
            return True


# date = "March 01, 2024"

# adjustments = [
#     [
#         # leave
#         [["jess"], [0, 0]],
#         [["sonaite"], ["12:00pm", "5:30pm"]],
#     ],
#     [
#         # programs and meetings
#         [["lea", "michelle"], ["9:15am", "10:15am"], "meeting"],
#         [["lea"], ["2:00pm", "3:00pm"], "meeting"],
#         [["rod"], ["6:00pm", "8:00pm"], "chess"],
#     ],
# ]

date = datetime.today().strftime("%B %d, %Y")
date = "March 01, 2024"
adjustments = [
    [
        [["cat"], [0, 0]], [["chris"], ["9:00am", "11:00pm"]]
    ],
    [
        # programs and meetings
    ]
]

daily_schedule = Schedule(date, adjustments)

# daily_schedule.change_doc()
